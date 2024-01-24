from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import Table
import os
import csv
import docx2txt
import re
import shutil
import sys

def create_section_directory(section_directory):
    for i in range(1, 21):
        command = f"mkdir {section_directory}section{i}"
        os.system(command)

def delete_sections_files(sections_directory):
    #Iterate through sections directory
    for filename in os.listdir(sections_directory):
        if filename.startswith("section"):
            #check if the directory is empty
            if os.listdir(sections_directory+filename):    
                #remove the files from the directory
                for file in os.listdir(sections_directory+filename):
                    os.remove(sections_directory+filename+"/"+file)
                
def analyze_word_document2(file_path):
    # load the word document
    doc = Document(file_path)
    element_order = []
    file_number = 0
    section_number = 0
    image_number = 1    

    code_block = ""
    old_style = ""
    new_style = ""

    for element in doc.element.body:
        if isinstance(element, CT_P):
            paragraph = Paragraph(element, None)
            paragraph = next((p for p in doc.paragraphs if p._element is element), None)
            
            if 'graphicData' in paragraph._p.xml:
                    if(element.text != ""):
                        #print(paragraph.style.name)
                        #print(element.text)
                        element_order.append((section_number, file_number, 'Paragraph', element.text, paragraph.style.name))
                        file_number += 1

                    element_order.append((section_number, file_number, 'Image', 'image'+str(image_number)+'.png', "none"))
                    file_number += 1
                    image_number += 1
            else:
                if(paragraph.text != ""):   
                    style = paragraph.style.name
                    is_bold = any(run.bold for run in paragraph.runs)
                    if(is_bold and paragraph.style.name == "Normal"):
                        style  = "Caption"

                
                    if style == "Heading 1":
                        section_number += 1
                        file_number = 1
                        element_order.append((section_number, file_number, 'Paragraph', paragraph.text, style))
                        file_number += 1

                        if code_block != "":
                            code_block = ""

                    elif style == "Code Standalone":
                        code_block = code_block + paragraph.text + "\n"

                    else: 
                        if code_block != "":
                            element_order.append((section_number, file_number, 'Paragraph', code_block, "Code Standalone"))
                            code_block = ""
                            file_number += 1
                        element_order.append((section_number, file_number, 'Paragraph', paragraph.text, style))
                        file_number += 1
                elif paragraph.text == "" and code_block != "":
                    element_order.append((section_number, file_number, 'Paragraph', code_block, "Code Standalone"))
                    code_block = ""
                    file_number += 1



        elif isinstance(element, CT_Tbl):
            # print("This is a table")
            #print(element)
            all_rows = []
            table = Table(element, None)

            for i, row in enumerate(table.rows):
                row_content = []

                for cell in row.cells:
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        if len(paragraph.text.strip()) > 0:
                            cell_text += paragraph.text

                    # Check if the cell_text contains a comma
                    if "," in cell_text:
                        # If yes, wrap the cell_text in quotes
                        cell_text = f'"{cell_text}"'

                    row_content.append(cell_text)

                # Append the row content to the list
                all_rows.append(row_content)

            element_order.append((section_number, file_number, 'Table', all_rows, "none"))

    return element_order    





def iterate_paragraphs(element_order,sections_directory,file_directory):
    doc = Document(file_directory)
    # counter = 0 
    table_counter = 1
    figure_counter = 1

    for section, number, element_type, content, style in element_order:
        if element_type == "Paragraph":
            if style == "Heading 1":
                output_file_path = f"{sections_directory}section{section}/section{section}title.txt"
            
            else:
                output_file_path = f"{sections_directory}section{section}/section{section}_{number}_output_text.txt"


            if style  == "Caption":
                if content.startswith("Fig"):
                    content = content.replace("Figure -. ", "Figure "+str(figure_counter)+". ")
                    content = content.replace("Figure - ", "Figure "+str(figure_counter)+". ")
                    
                    figure_counter += 1
                elif content.startswith("Table"):
                    content = content.replace("Table .. ", "Table "+str(table_counter)+". ")
                    content = content.replace("Table . ", "Table "+str(table_counter)+". ")
                    table_counter += 1
            if content.startswith("Figure") and style != "Caption":
                content = content.replace("Figure -. ", "Figure "+str(figure_counter)+". ")
                content = content.replace("Figure - ", "Figure "+str(figure_counter)+". ")
                figure_counter += 1
                style = "Caption"

            if content.startswith("Table") and style != "Caption":
                content = content.replace("Table .. ", "Table "+str(table_counter)+". ")
                content = content.replace("Table . ", "Table "+str(table_counter)+". ")
                table_counter += 1
                style = "Caption"

            if style == "Code Standalone":
                with open(output_file_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(content)
            else:    
                with open(output_file_path, 'a', encoding='utf-8') as txt_file:
                    txt_file.write(f"Style: {style}\nContent: {content}\n\n")
            
            


def iterate_tables(element_order,sections_directory,file_directory):
    for section, number, element_type, content, style in element_order:
        if element_type == "Table":
            output_file_path = f"{sections_directory}section{section}/section{section}_{number}_output_table.csv"
            with open(output_file_path, 'w', newline='', encoding='utf-8') as csv_file:
                csv_writer = csv.writer(csv_file)

                for row in content:
                    cleaned_row = [cell.replace('\u200b', '') for cell in row]
                    csv_writer.writerow(cleaned_row)


def iterate_images(element_order,sections_directory,file_directory):
    image_directory = f"{sections_directory}images/"
    docx2txt.process(file_directory, image_directory)

    for section, number, element_type, content, style in element_order:
        if(element_type == "Image"):   
            original_file_path = f"{image_directory}{content}"
            new_file_path = f"{sections_directory}section{section}/section{section}_{number}_output_image.png"

            if os.path.exists(original_file_path):
                os.rename(original_file_path, new_file_path)
                




def main(customer_name):
    word_document = "initial.docx"
    output_directory = f"../documentation/{customer_name}/"
    file_directory = f"../documentation/{customer_name}/{word_document}"
    sections_directory = f"../documentation/{customer_name}/public/"

    delete_sections_files(sections_directory)
    element_order=analyze_word_document2(file_directory)
    # element_order = analyze_word_document(file_directory)
    #for element in element_order:
    #    print(element)
    iterate_paragraphs(element_order, sections_directory, file_directory)
    iterate_tables(element_order, sections_directory, file_directory)
    iterate_images(element_order, sections_directory, file_directory)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract.py <customer_name>")
        sys.exit(1)

    customer_name = sys.argv[1]
    main(customer_name)


